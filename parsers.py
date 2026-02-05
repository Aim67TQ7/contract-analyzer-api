"""
Document parsers for contract text extraction.
Supports PDF, DOCX, and plain text files.
"""

import io
from typing import Optional


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def parse_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return parse_docx(file_bytes)
    elif ext in ("txt", "text", "md", "csv"):
        return parse_txt(file_bytes)
    else:
        # Try as plain text
        return parse_txt(file_bytes)
